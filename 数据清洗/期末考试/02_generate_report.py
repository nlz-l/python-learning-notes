from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
OUT_DIR = BASE_DIR / "output"
FIG_DIR = OUT_DIR / "figures"
TEMPLATE_PATH = BASE_DIR / "2023109223-刘文博-数据清洗课程报告.docx"
REPORT_PATH = BASE_DIR / "2023109223-刘文博-数据清洗课程报告_python生成.docx"


def set_font(run, name: str = "宋体", size: int = 12, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_base(paragraph, first_indent: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_indent else None


def add_level1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_base(p)
    run = p.add_run(text)
    set_font(run, "黑体", 14, True)


def add_level2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_base(p)
    run = p.add_run(text)
    set_font(run, "黑体", 12, True)


def add_body(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    set_paragraph_base(p, first_indent=indent)
    run = p.add_run(text)
    set_font(run, "宋体", 12, False)


def add_blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_base(p)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def style_table(table) -> None:
    table.style = "Normal Table"
    set_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_base(p)
                for run in p.runs:
                    set_font(run, "宋体", 10, None)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base(p)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_font(run, "宋体", 12, False)


def add_figure(doc: Document, image_name: str, caption: str) -> None:
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.35))
    add_caption(doc, caption)


def remove_template_body(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    start_index = None
    for idx, child in enumerate(children):
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if text.startswith("1数据清洗综合应用实践过程展示"):
            start_index = idx
            break
    if start_index is None:
        raise RuntimeError("未在模板中找到正文起始位置。")
    for child in children[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def read_summary() -> tuple[dict, list[dict[str, str]]]:
    with open(OUT_DIR / "summary.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    return data["quality_summary"], data["cleaning_log"]


def add_feature_table(doc: Document) -> None:
    rows = [
        ["date", "月份日期", "年月", "日期"],
        ["production_10k", "汽车月度产量", "万辆", "浮点数"],
        ["production_missing_flag", "原始月份是否缺失", "0/1", "整数"],
        ["production_filled_10k", "清洗后汽车月度产量", "万辆", "浮点数"],
        ["production_yoy", "汽车月度产量同比增速", "比例", "浮点数"],
        ["nev_sales_10k", "新能源汽车年度销量", "万辆", "浮点数"],
        ["china_world_share_pct", "中国电动汽车销量全球占比", "%", "浮点数"],
        ["policy_stage", "新能源汽车发展阶段", "类别", "文本"],
    ]
    add_caption(doc, "表1-1  中国新能源汽车数据集主要字段说明")
    table = doc.add_table(rows=1 + len(rows), cols=4)
    headers = ["特征名称", "含义", "单位", "数据类型"]
    for j, header in enumerate(headers):
        table.cell(0, j).text = header
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    style_table(table)


def add_cleaning_log_table(doc: Document, log: list[dict[str, str]]) -> None:
    add_caption(doc, "表1-2  数据清洗过程记录")
    table = doc.add_table(rows=1 + len(log), cols=2)
    table.cell(0, 0).text = "清洗步骤"
    table.cell(0, 1).text = "处理说明"
    for i, item in enumerate(log, start=1):
        table.cell(i, 0).text = item["step"]
        table.cell(i, 1).text = item["detail"]
    style_table(table)


def add_references(doc: Document) -> None:
    refs = [
        "[1] 郭志懋, 周傲英. 数据质量和数据清洗研究综述[J]. 软件学报, 2002, 13(11): 2076-2082.",
        "[2] 杨辅祥, 刘云超, 段智华. 数据清理综述[J]. 计算机应用研究, 2002, 19(3): 3-5.",
        "[3] 王曰芬, 章成志, 张蓓蓓, 等. 数据清洗研究综述[J]. 现代图书情报技术, 2007(12): 50-56.",
        "[4] 赵一凡, 卞良, 丛昕. 数据清洗方法研究综述[J]. 软件导刊, 2017, 16(12): 222-224.",
        "[5] 郝爽, 李国良, 冯建华, 等. 结构化数据清洗方法综述[J]. 清华大学学报(自然科学版), 2018, 58(12): 1037-1050.",
        "[6] 王铭军, 潘巧明, 刘真, 陈为. 可视数据清洗综述[J]. 中国图象图形学报, 2015, 20(4): 468-482.",
        "[7] 杨东华, 李宁宁, 王宏志, 等. 基于任务合并的并行大数据清洗过程优化[J]. 计算机学报, 2016, 39(1): 97-108.",
        "[8] 国家市场监督管理总局, 中国国家标准化管理委员会. 信息技术 数据质量评价指标: GB/T 36344-2018[S]. 北京: 中国标准出版社, 2018.",
        "[9] 国务院办公厅. 新能源汽车产业发展规划(2021-2035年)[Z]. 2020.",
        "[10] China Data Portal. China NEV Sales, China Monthly Car Production, Electric Vehicle Sales: China vs World[DS/OL]. 2026.",
    ]
    add_level2(doc, "参考文献：")
    for ref in refs:
        add_body(doc, ref, indent=False)


def build_report() -> None:
    summary, log = read_summary()
    doc = Document(TEMPLATE_PATH)
    remove_template_body(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    add_level1(doc, "1数据清洗综合应用实践过程展示")
    add_level2(doc, "1.1案例题目")
    add_body(doc, "（1）案例题目：基于 Python 的中国新能源汽车数据清洗与趋势分析。", indent=False)
    add_body(doc, "（2）数据集来源说明：本案例数据来自 China Data Portal 提供的公开 API，其中新能源汽车年度销量来源标注为中国汽车工业协会（CAAM），中国汽车月度产量来源标注为国家统计局（NBS），中国与全球电动汽车销量对比来源标注为 IEA、EV-Volumes、中国汽车工业协会等公开统计口径。数据接口包括：https://chinadata.live/api/v2/data/china-nev-sales、https://chinadata.live/api/v2/data/china-car-production-monthly、https://chinadata.live/api/v2/data/ev-sales-china-vs-world。", indent=False)
    add_body(doc, "（3）案例应用背景说明：新能源汽车是我国汽车产业转型升级的重要方向，也是当前制造业、低碳交通和智能网联汽车领域的热点。本案例通过清洗新能源汽车销量、汽车产量和全球对比数据，分析中国新能源汽车市场增长趋势，并展示时间序列数据清洗的基本流程。", indent=False)

    add_level2(doc, "1.2 分析数据集")
    add_body(doc, "（1）原始数据集的名称：中国新能源汽车年度销量数据、中国汽车月度产量数据、中国与全球电动汽车销量对比数据。", indent=False)
    add_body(doc, f"（2）数据的记录数量：月度汽车产量原始数据 {summary['monthly_rows_raw']} 条，补齐完整月份后为 {summary['monthly_rows_clean']} 条；年度新能源汽车分析表为 {summary['annual_rows_clean']} 条。", indent=False)
    add_body(doc, "（3）数据的字段数量：清洗后的月度产量表包含 11 个字段，年度新能源汽车分析表包含 13 个字段。", indent=False)
    add_body(doc, "（4）数据的字段名称及内涵介绍如下表所示。", indent=False)
    add_feature_table(doc)
    add_body(doc, f"（5）数据集存在的问题：原始 API 表格本身没有空单元格，但月度产量数据不是完整月份序列。按完整月份索引补齐后，可以发现 {summary['inserted_missing_months']} 个未单独报告月份，这类问题属于时间序列缺失；不同数据集的时间频率和单位不同，需要统一年份字段、单位口径和字段命名；月度产量存在少量极端值，需要用统计方法检查；年度销量与月度产量汇总需要按年份合并。", indent=False)
    add_body(doc, "（6）清洗目标：统一字段格式，补齐完整月份索引，标记并填补未报告月份，检查并处理异常值，构造同比增速、政策阶段、全球占比等分析字段，输出可复现的清洗后数据、质量汇总表和图表。", indent=False)

    add_level2(doc, "1.3 数据清洗过程")
    add_body(doc, "针对原始数据存在的时间序列不完整、字段命名不统一、单位口径不同、缺失月份和异常值等问题，本报告使用 Python 完成数据清洗。具体过程如下表所示。")
    add_cleaning_log_table(doc, log)

    add_level2(doc, "1.4 数据清洗前后对比")
    add_body(doc, f"本节只展示数据清洗前后的直接对比。清洗后，月度产量表和年度分析表的缺失值均为 0。需要说明的是，原始 API 数据没有空单元格，真正的问题是缺少部分月份记录；月度数据从原始 {summary['monthly_rows_raw']} 条补齐为 {summary['monthly_rows_clean']} 条，填补前识别出 {summary['missing_before_total']} 个时间序列缺失值，清洗后全部补齐，解决了时间序列不连续的问题。异常值方面，月度产量使用 IQR 方法检查，截尾范围为 [{summary['production_iqr_low']}, {summary['production_iqr_high']}]，共调整 {summary['production_outliers']} 个极端值。")
    add_figure(doc, "01_missing_before_after.png", "图1-1  补齐月份后填补前与清洗后缺失值对比")
    add_figure(doc, "05_production_boxplot_before_after.png", "图1-2  月度产量异常值处理前后对比")

    add_level2(doc, "1.5 清洗后数据分析结果")
    add_body(doc, f"本节基于清洗后的数据进行趋势分析和业务对比分析，以下图表不属于清洗前后对比图。年度表将新能源汽车销量、年度汽车产量汇总和全球对比数据按年份合并，便于进行趋势分析。2024 年中国新能源汽车销量为 {summary['latest_nev_sales_10k']} 万辆，折合约 {summary['latest_nev_sales_million']} 百万辆；中国电动汽车销量全球占比约为 {summary['latest_china_world_share_pct']}%。从趋势看，2021 年以后中国新能源汽车销量进入规模化增长阶段，全球占比长期保持较高水平。")
    add_figure(doc, "02_nev_sales_trend.png", "图1-3  清洗后中国新能源汽车年度销量趋势")
    add_figure(doc, "03_nev_yoy_growth.png", "图1-4  清洗后中国新能源汽车年度销量同比增速")
    add_figure(doc, "04_monthly_car_production.png", "图1-5  清洗后中国汽车月度产量趋势")
    add_figure(doc, "06_china_world_share.png", "图1-6  清洗后中国电动汽车销量全球占比")
    add_figure(doc, "07_china_vs_world_ev_sales.png", "图1-7  清洗后中国与世界其他地区电动汽车销量业务对比")
    add_figure(doc, "08_annual_correlation_heatmap.png", "图1-8  清洗后年度新能源汽车关键指标相关系数热力图")

    add_blank(doc)
    add_level1(doc, "2 数据清洗综述")
    add_level2(doc, "2.1基本概念及内涵")
    add_body(doc, "数据清洗是数据预处理的重要组成部分，主要用于识别并处理数据中的缺失值、重复值、异常值、格式不一致、编码不统一和逻辑冲突等问题。国内关于数据质量和数据清洗的研究普遍认为，数据清洗不仅是简单地删除错误记录，更重要的是结合业务背景判断问题产生的原因，并通过可复现的规则对数据进行修正和验证。对于本报告使用的中国新能源汽车数据而言，数据质量问题主要体现在时间序列不完整、月度统计缺口、指标单位不一致和跨数据源合并口径不同等方面。")

    add_level2(doc, "2.2 数据清洗主要内容。")
    add_body(doc, "数据清洗的主要内容包括数据剖析、缺失值处理、重复值处理、异常值检测、格式标准化、数据转换、数据集成和清洗结果验证。常用方法包括众数填补、中位数填补、分组填补、插值法、IQR 异常值检测、标准差法、正则表达式解析和基于业务规则的逻辑校验等。Python 中的 Pandas、NumPy、Matplotlib 和 Seaborn 能够较好地支持课程案例中的数据清洗与可视化。")
    add_body(doc, "在本案例中，月度产量数据首先被转换为标准日期格式，然后按完整月份序列补齐。对未单独报告的月份，本报告保留 production_missing_flag 字段作为缺失标记，并用同月历史中位数填补。对于极端月度产量，采用 IQR 方法进行检查和截尾处理。年度数据部分则按 year 字段合并新能源汽车销量、汽车产量和全球对比数据，并构造同比增速、全球占比和政策阶段等分析字段。")

    add_level2(doc, "2.3 应用领域")
    add_body(doc, "数据清洗广泛应用于政府统计、金融风控、医疗健康、电子商务、工业制造、交通出行和能源管理等领域。在新能源汽车领域，数据清洗可以用于产销量统计、充电基础设施分析、用户行为研究、产业链监测、市场预测和政策效果评估等任务。高质量的数据清洗能够提升后续统计分析和预测建模的可靠性。")

    add_level2(doc, "2.4 热点问题或难点问题")
    add_body(doc, "当前数据清洗的热点问题包括大规模数据自动清洗、多源异构数据融合、缺失值机制识别、异常值自动检测、清洗规则可解释性和清洗过程可追溯等。新能源汽车数据具有统计口径多、更新频率不同、产业链环节复杂等特点，因此在清洗过程中既要关注数值分布，也要关注数据来源和指标含义。例如，新能源汽车销量和汽车产量属于不同统计口径，不能简单把二者相除解释为严格渗透率，只能作为口径对比指标辅助观察趋势。")

    add_level2(doc, "2.5 今后的发展趋势")
    add_body(doc, "未来数据清洗将更加重视自动化、智能化和可解释性。一方面，机器学习和大语言模型可以辅助识别字段语义、推荐清洗规则和生成数据质量报告；另一方面，关键行业数据仍然需要人工结合业务背景进行判断。对于新能源汽车这类政策和市场共同驱动的热点数据，后续可以继续融合充电桩、地区经济、人口规模和政策补贴等数据，构建更加完整的数据分析体系。")

    add_references(doc)
    try:
        doc.save(REPORT_PATH)
        print(f"报告已按模板格式生成: {REPORT_PATH}")
    except PermissionError:
        fallback_path = REPORT_PATH.with_name(REPORT_PATH.stem + "_新版.docx")
        doc.save(fallback_path)
        print(f"目标报告被占用，已另存为: {fallback_path}")


if __name__ == "__main__":
    build_report()
